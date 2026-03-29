import html2text
from docx import Document
from docx.shared import Inches
import re
import config
from bs4 import BeautifulSoup
import logging

class DocumentConverter:
    def __init__(self):
        self.h2t = html2text.HTML2Text()
        self.h2t.ignore_links = False
        self.h2t.ignore_images = False
        self.h2t.body_width = 0
        self.h2t.unicode_snob = True
        self.h2t.escape_snob = False  

    def html_to_markdown(self, html_content):
        """Convert HTML to Markdown while preserving formatting, avoiding duplication in list items."""
        soup = BeautifulSoup(html_content, 'html.parser')

        # Remove <p> tags inside <li> tags to prevent duplication
        for li in soup.find_all('li'):
            for p in li.find_all('p', recursive=False):
                p.unwrap()

        markdown_content = self.h2t.handle(str(soup))
        markdown_content = re.sub(r'\n\s*\n\s*\n', '\n\n', markdown_content)
        return markdown_content.strip()

    def clear_document_content(self, doc):
        """Clear all content from document except header/footer."""
        for para in doc.paragraphs:
            p_element = para._element
            p_element.getparent().remove(p_element)
        for table in doc.tables:
            tbl_element = table._element
            tbl_element.getparent().remove(tbl_element)

    def add_formatted_text(self, paragraph, text):
        """Add text with markdown formatting to a paragraph."""
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Consolas'
            elif part:
                paragraph.add_run(part)

    def convert_and_append(self, html_content, docx_path):
        """Convert HTML and append to DOCX (clears previous content)."""
        markdown_content = self.html_to_markdown(html_content)

        try:
            doc = Document(docx_path)
            self.clear_document_content(doc)
        except Exception:
            doc = Document()

        lines = markdown_content.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i].strip()

            # Headers
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('# ').strip()
                level = min(level, 9)
                doc.add_heading(text, level=level)
                i += 1

            # Lists
            elif line.lstrip().startswith(('* ', '- ')) or re.match(r'^\s*\d+\.\s', line.lstrip()):
                while i < len(lines):
                    list_line = lines[i].lstrip()
                    is_bullet = list_line.startswith(('* ', '- '))
                    is_numbered = re.match(r'^\d+\.\s', list_line)

                    if not (is_bullet or is_numbered):
                        break

                    if is_bullet:
                        bullet_symbol = "• "
                        text = bullet_symbol + list_line[2:]
                    elif is_numbered:
                        number = re.match(r'^(\d+)\.\s', list_line).group(1)
                        text = f"{number}. " + re.sub(r'^\d+\.\s', '', list_line)

                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25)
                    self.add_formatted_text(p, text)
                    i += 1

            # Code blocks
            elif line.startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1

                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'

            # Blockquotes
            elif line.startswith('>'):
                quote_text = line.lstrip('> ').strip()
                try:
                    p = doc.add_paragraph(style='Quote')
                except KeyError:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)

                run = p.add_run(quote_text)
                run.italic = True
                i += 1

            # Regular text or blank line
            else:
                if line:
                    p = doc.add_paragraph()
                    self.add_formatted_text(p, line)
                else:
                    doc.add_paragraph()
                i += 1

        doc.save(docx_path)
        logging.info(f"Successfully converted HTML and updated {docx_path}")

def convert_html_to_docx(html_content: str, docx_path: str) -> str:
    """Convert HTML content to DOCX format and save it."""
    converter = DocumentConverter()
    converter.convert_and_append(html_content, docx_path)
    return docx_path