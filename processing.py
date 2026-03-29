#Pending points:
# 1. Implement the deletion of the initial local file uploads after processing.
# 2. Ensure that the chatbot session can be properly initialized and messages can be sent.
# 3. Add .env inside the docker file only

import os 
from werkzeug.utils import secure_filename
import PyPDF2
import docx
from app_files.api.l1_transcpit_generation import process_videos, process_audio
from app_files.api.l2_direct_html_brd import ask_brd_gen
from app_files.api.l3_brd_qa import ask_brd_qa
from app_files.api.l5_brd_gen import brd_gen
from app_files.chatbot.chatbot_api import initialize_chatbot_session, send_message_to_chatbot
from app_files.services.html_to_docx import convert_html_to_docx
from app_files.services.link_to_video import download_meeting_video
from typing import List, Dict, Any
from app_files.chatbot.file_manager import FileStorageManager
import config
import tempfile
import shutil
import sys
import logging

logging.basicConfig(level=logging.INFO,format='%(asctime)s - %(levelname)s - %(message)s',  handlers=[logging.StreamHandler(sys.stdout)])

storage = config.storage_type

ALLOWED_EXTENSIONS = config.file_type

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

file_manager = FileStorageManager(storage)

def extraction_files(request,app, unique_id,user_id):
    logging.info(f"Starting file extraction for user_id: {user_id}, unique_id: {unique_id}")
    content = ''
    transcript= ''

    if 'video_path' not in request.form:
        gcs_paths = []
    else:
        gcs_paths = request.form.getlist('video_path') 
    logging.info(f"Received GCS paths: {gcs_paths}")
    for gcs_path in gcs_paths:
        # Extract filename from GCS path
        filename = os.path.basename(gcs_path)
        logging.info(f"Processing file: {filename} from GCS path: {gcs_path}")
        file_extension = os.path.splitext(filename)[1].lower()
        logging.info(f"Processing file: {filename} from GCS path: {gcs_path}")

        temp_path = os.path.join("/tmp", f"{filename}")

        # Read file content from GCS
        file_manager.read_file_from_path(gcs_path, temp_path)

        #Read the video file in bytes from temp path
        with open(temp_path, 'rb') as f:
            video_content = f.read()

        # === Handle file type processing ===
        if file_extension == '.pdf':
            with open(temp_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                content = ''
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    text = page.extract_text() or ''
                    print(f"Page {page_num + 1}:")
                    print(text)
                    content += text + '\n'

        elif file_extension == '.docx':
            doc = docx.Document(temp_path)
            content = '\n'.join([para.text for para in doc.paragraphs])

        elif file_extension == '.txt':
            with open(temp_path, 'r', encoding='utf-8') as f:
                content = f.read()

        elif file_extension == '.mp4':
            logging.info(f"Processing video file: {filename}")
            transcript += process_videos([temp_path])

        elif file_extension == '.mp3':
            logging.info(f"Processing audio file: {filename}")
            transcript += process_audio(temp_path)

        else:
            print(f"Invalid file format: {filename}")
            continue

        # Clean up temp file
        os.remove(temp_path)
        file_manager.delete_file_from_path(gcs_path)

    logging.info("Completed processing GCS files.")

    # check if files are uploaded
    if 'files' not in request.files:
        files = []
    else:
        files = request.files.getlist('files')
    for file in files: 
        if file.filename == '':
            return 'No selected file', 400
        
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file_extension = os.path.splitext(filename)[1].lower()
            
            if(storage=='local'):
                # Save the file to a temporary location
                file_path=file_manager.save_file(user_id, unique_id, filename, file.read())
            elif(storage=='gcs'):
                file_path = os.path.join("/tmp", filename)
                
            file.save(file_path)
            
            
            if file_extension == '.pdf':
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(file)
                    
                    # Get the number of pages in the PDF
                    num_pages = len(reader.pages)
                    content = ''
                    
                    # Loop through all pages and extract text
                    for page_num in range(num_pages):
                        page = reader.pages[page_num]
                        text = page.extract_text()
                        print(f"Page {page_num + 1}:")
                        print(text)
                        print("\n" + "="*50 + "\n")  # Divider between pages
                        content = content + text

                    

            elif file_extension == '.docx':
                doc = docx.Document(file_path)
                content += '\n'.join([para.text for para in doc.paragraphs])
            elif file_extension == '.txt':
                with open(file_path, 'r') as f:
                    content += f.read()
            elif file_extension == '.mp4':
                # Handle video files if needed
                transcript += process_videos([file_path])
            elif file_extension == '.mp3':
                transcript += process_audio(file_path)
            else:
                return 'Invalid file format', 400

            os.remove(file_path)
    
    logging.info("Completed processing uploaded files.")

    if content != '':
        file_manager.save_file(user_id,unique_id,'li_content.txt',content)
    if transcript != '':
        file_manager.save_file(user_id,unique_id,'l1_transcript.txt',transcript)


    if not content and not transcript:
        return 'No content extracted from the file', 400
    
    return transcript, content


def intial_upload(unique_id,user_id):

    #Read the transcript from the file
    try:
        transcript=file_manager.read_file(user_id, unique_id, 'l1_transcript.txt').decode('utf-8')
    except Exception as e:
        print(f"Error reading transcript file: {e}")
        transcript = ''
    try:
        # Read the content from the file
        content = file_manager.read_file(user_id, unique_id, 'li_content.txt').decode('utf-8')
    except Exception as e:
        print(f"Error reading content file: {e}")
        content = ''


    temp_brd = ask_brd_gen(transcript+"\n\n Following is data of supporting documents : \n"+content)
    
    #save the BRD.html
    file_manager.save_file(user_id, unique_id, 'brd.html', temp_brd.encode('utf-8'))

    return temp_brd

# generate questions based on the BRD content analysis
def brd_qa(user_id, unique_id):
    
    try:
        brd_content= file_manager.read_file(user_id, unique_id, 'brd.html').decode('utf-8')
    except Exception as e:
        print(f"Error reading BRD file: {e}")
        return None

    # Generate questions based on BRD analysis
    temp_questions = ask_brd_qa(brd_content)

    # Save the questions to a JSON file
    file_manager.save_json_file(user_id, unique_id, 'brd_questions.json', temp_questions)

    return temp_questions

#CHATBOT
def chatbot_initialization(user_id, unique_id: str) -> Dict[str, Any]:
    """
    Initialize a chatbot session for frontend use.

    Args:
        unique_id: Unique identifier for the chatbot session
        chatbot_history: Optional chat history. If None, defaults to [{"type": "human", "content": "hi bhavya"}]
    
    Returns:
        Dict containing:
        - status: "success" or "error"
        - bot_response: The bot's response message
        - session_id: The session identifier
        - error: Error message (if status is "error")
    """
    return initialize_chatbot_session(user_id, unique_id)

def chatbot_conversation(user_id, unique_id: str, user_message: str) -> Dict[str, Any]:
    """
    Continue a chatbot conversation.
    
    Args:
        unique_id: Unique identifier for the chatbot session
        user_message: The user's message/reply
        chatbot_history: Current conversation history in the format:
            [
                {"type": "human", "content": "user message"},
                {"type": "ai", "content": "bot response"},
                ...
            ]
    
    Returns:
        Dict containing:
        - status: "success" or "error"
        - bot_response: The bot's response message
        - session_id: The session identifier
        - error: Error message (if status is "error")
    """
    return send_message_to_chatbot(user_id, unique_id, user_message)




# GENERATE BRD and Uploading temporary files
def generate_brd(userid, unique_id, content):

    try:
        brd_content = file_manager.read_file(userid, unique_id, 'brd.html').decode('utf-8') if file_manager.read_file(userid, unique_id, 'brd.html') else ''
    except Exception as e:
        print(f"Error reading BRD file: {e}")
        brd_content = ''
    try:
        chatbot_data = file_manager.read_json_file(userid, unique_id, 'chatbot.json') if file_manager.read_json_file(userid, unique_id, 'chatbot.json') else {}
    except Exception as e:
        print(f"Error reading chatbot data file: {e}")
        chatbot_data = {}
    result = brd_gen(unique_id, brd_content, chatbot_data, content)

    # Save the generated BRD to a file
    brd_file_path = file_manager.save_file(userid, unique_id, 'brd.html', result.encode('utf-8'))
    return result

def download_brd(userid, unique_id, filename, content):
    header_doc = config.brd_document_path
    temp_dir = tempfile.gettempdir()
    new_filename = f"{filename}.docx"
    temp_path = os.path.join(temp_dir, new_filename)
    shutil.copyfile(header_doc, temp_path)
    convert_html_to_docx(content, temp_path)
    return temp_path

def download_video(user_id, unique_id, token, meet_url, meet_name):
    return download_meeting_video(user_id, unique_id, token, meet_url, meet_name)

def session_delete(user_id, unique_id):
    """
    Delete a session and its associated files from local or GCS.
    Returns (status_code, message).
    """
    try:
        deleted = file_manager.delete_session(user_id, unique_id)
        if not deleted:
            return 404, f"No session data found for user {user_id}, session {unique_id}"

        logging.info(f"Session {unique_id} for user {user_id} deleted successfully.")
        return 200, "Session deleted successfully"
    except Exception as e:
        logging.error(f"Error deleting session {unique_id} for user {user_id}: {e}")
        return 500, str(e)